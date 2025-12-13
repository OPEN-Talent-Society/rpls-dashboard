#!/usr/bin/env python3
"""
Universal document text extraction utility.
Supports: PDF, DOCX, DOC, XLSX, XLS, CSV, TXT, HTML, HTM, MD, JSON, XML

Usage:
    python3 extract-document-text.py <file_path>

Output: JSON with extracted text and metadata
    {
        "success": true,
        "file_path": "...",
        "file_type": "pdf",
        "text": "extracted content...",
        "metadata": {
            "pages": 5,
            "word_count": 1234,
            ...
        },
        "error": null
    }
"""

import sys
import os
import json
import subprocess
from pathlib import Path

def extract_pdf(file_path):
    """Extract text from PDF using pdftotext (poppler)"""
    try:
        result = subprocess.run(
            ['pdftotext', '-layout', file_path, '-'],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            text = result.stdout
            # Count pages by page breaks
            pages = text.count('\x0c') + 1 if text else 0
            return {
                'text': text,
                'metadata': {'pages': pages, 'extraction_method': 'pdftotext'}
            }
    except FileNotFoundError:
        pass
    except Exception as e:
        return {'error': str(e)}

    # Fallback: try PyPDF2 if available
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = '\n'.join(page.extract_text() or '' for page in reader.pages)
            return {
                'text': text,
                'metadata': {'pages': len(reader.pages), 'extraction_method': 'PyPDF2'}
            }
    except ImportError:
        pass
    except Exception as e:
        return {'error': f'PDF extraction failed: {e}'}

    return {'error': 'No PDF extraction tool available (install pdftotext or PyPDF2)'}

def extract_docx(file_path):
    """Extract text from DOCX using python-docx"""
    try:
        from docx import Document
        doc = Document(file_path)

        paragraphs = [p.text for p in doc.paragraphs]

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)

        text = '\n'.join(paragraphs)
        return {
            'text': text,
            'metadata': {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'extraction_method': 'python-docx'
            }
        }
    except ImportError:
        pass
    except Exception as e:
        return {'error': str(e)}

    # Fallback: try pandoc
    try:
        result = subprocess.run(
            ['pandoc', '-f', 'docx', '-t', 'plain', file_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            return {
                'text': result.stdout,
                'metadata': {'extraction_method': 'pandoc'}
            }
    except FileNotFoundError:
        pass
    except Exception as e:
        return {'error': str(e)}

    return {'error': 'No DOCX extraction tool available'}

def extract_doc(file_path):
    """Extract text from legacy DOC using textutil (macOS) or antiword"""
    # macOS textutil
    try:
        result = subprocess.run(
            ['textutil', '-convert', 'txt', '-stdout', file_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            return {
                'text': result.stdout,
                'metadata': {'extraction_method': 'textutil'}
            }
    except FileNotFoundError:
        pass
    except Exception as e:
        pass

    # Try pandoc as fallback
    try:
        result = subprocess.run(
            ['pandoc', '-f', 'doc', '-t', 'plain', file_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            return {
                'text': result.stdout,
                'metadata': {'extraction_method': 'pandoc'}
            }
    except:
        pass

    return {'error': 'No DOC extraction tool available (install textutil or pandoc)'}

def extract_xlsx(file_path):
    """Extract text from XLSX using openpyxl"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)

        all_text = []
        sheet_info = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_text = [str(cell) if cell is not None else '' for cell in row]
                if any(row_text):
                    rows.append(' | '.join(row_text))

            if rows:
                all_text.append(f"## Sheet: {sheet_name}\n" + '\n'.join(rows))
                sheet_info.append({'name': sheet_name, 'rows': len(rows)})

        wb.close()
        return {
            'text': '\n\n'.join(all_text),
            'metadata': {
                'sheets': sheet_info,
                'total_sheets': len(wb.sheetnames),
                'extraction_method': 'openpyxl'
            }
        }
    except ImportError:
        return {'error': 'openpyxl not installed'}
    except Exception as e:
        return {'error': str(e)}

def extract_xls(file_path):
    """Extract text from legacy XLS using xlrd or pandas"""
    try:
        import pandas as pd
        # Read all sheets
        xlsx = pd.ExcelFile(file_path)
        all_text = []
        sheet_info = []

        for sheet_name in xlsx.sheet_names:
            df = pd.read_excel(xlsx, sheet_name=sheet_name)
            text = df.to_string()
            all_text.append(f"## Sheet: {sheet_name}\n{text}")
            sheet_info.append({'name': sheet_name, 'rows': len(df)})

        return {
            'text': '\n\n'.join(all_text),
            'metadata': {
                'sheets': sheet_info,
                'extraction_method': 'pandas'
            }
        }
    except ImportError:
        return {'error': 'pandas not installed'}
    except Exception as e:
        return {'error': str(e)}

def extract_csv(file_path):
    """Extract text from CSV"""
    try:
        import csv
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            reader = csv.reader(f)
            rows = [' | '.join(row) for row in reader]

        return {
            'text': '\n'.join(rows),
            'metadata': {
                'rows': len(rows),
                'extraction_method': 'csv'
            }
        }
    except Exception as e:
        return {'error': str(e)}

def extract_html(file_path):
    """Extract text from HTML using BeautifulSoup"""
    try:
        from bs4 import BeautifulSoup
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')

        # Remove script and style elements
        for script in soup(['script', 'style']):
            script.decompose()

        text = soup.get_text(separator='\n', strip=True)

        # Get title if present
        title = soup.title.string if soup.title else None

        return {
            'text': text,
            'metadata': {
                'title': title,
                'extraction_method': 'beautifulsoup'
            }
        }
    except ImportError:
        # Fallback to simple regex
        try:
            import re
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            text = re.sub(r'<[^>]+>', ' ', content)
            text = re.sub(r'\s+', ' ', text)
            return {
                'text': text.strip(),
                'metadata': {'extraction_method': 'regex'}
            }
        except Exception as e:
            return {'error': str(e)}
    except Exception as e:
        return {'error': str(e)}

def extract_json(file_path):
    """Extract text from JSON files"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        def flatten_json(obj, prefix=''):
            items = []
            if isinstance(obj, dict):
                for k, v in obj.items():
                    new_key = f"{prefix}.{k}" if prefix else k
                    items.extend(flatten_json(v, new_key))
            elif isinstance(obj, list):
                for i, v in enumerate(obj):
                    items.extend(flatten_json(v, f"{prefix}[{i}]"))
            else:
                items.append(f"{prefix}: {obj}")
            return items

        text = '\n'.join(flatten_json(data))
        return {
            'text': text,
            'metadata': {
                'type': type(data).__name__,
                'extraction_method': 'json'
            }
        }
    except Exception as e:
        return {'error': str(e)}

def extract_xml(file_path):
    """Extract text from XML files"""
    try:
        import xml.etree.ElementTree as ET
        tree = ET.parse(file_path)
        root = tree.getroot()

        def get_text(elem):
            text_parts = []
            if elem.text:
                text_parts.append(elem.text.strip())
            for child in elem:
                text_parts.extend(get_text(child))
            if elem.tail:
                text_parts.append(elem.tail.strip())
            return text_parts

        text = '\n'.join(filter(None, get_text(root)))
        return {
            'text': text,
            'metadata': {
                'root_tag': root.tag,
                'extraction_method': 'xml.etree'
            }
        }
    except Exception as e:
        return {'error': str(e)}

def extract_text(file_path):
    """Extract text from plain text files"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read()
        return {
            'text': text,
            'metadata': {'extraction_method': 'plain'}
        }
    except Exception as e:
        return {'error': str(e)}

def extract_document(file_path):
    """Main extraction function - routes to appropriate extractor"""
    path = Path(file_path)

    if not path.exists():
        return {
            'success': False,
            'file_path': file_path,
            'file_type': None,
            'text': None,
            'metadata': None,
            'error': 'File not found'
        }

    ext = path.suffix.lower()
    file_type = ext.lstrip('.')

    extractors = {
        '.pdf': ('pdf', extract_pdf),
        '.docx': ('docx', extract_docx),
        '.doc': ('doc', extract_doc),
        '.xlsx': ('xlsx', extract_xlsx),
        '.xls': ('xls', extract_xls),
        '.csv': ('csv', extract_csv),
        '.html': ('html', extract_html),
        '.htm': ('html', extract_html),
        '.json': ('json', extract_json),
        '.xml': ('xml', extract_xml),
        '.txt': ('txt', extract_text),
        '.md': ('md', extract_text),
        '.markdown': ('md', extract_text),
        '.rst': ('rst', extract_text),
        '.rtf': ('rtf', extract_text),  # Will try pandoc
        '.css': ('css', extract_text),
        '.js': ('js', extract_text),
        '.py': ('py', extract_text),
        '.sh': ('sh', extract_text),
        '.yaml': ('yaml', extract_text),
        '.yml': ('yml', extract_text),
        '.toml': ('toml', extract_text),
        '.ini': ('ini', extract_text),
        '.conf': ('conf', extract_text),
        '.log': ('log', extract_text),
    }

    if ext in extractors:
        file_type, extractor = extractors[ext]
        result = extractor(file_path)
    else:
        # Try as plain text for unknown extensions
        result = extract_text(file_path)
        file_type = 'unknown'

    if 'error' in result:
        return {
            'success': False,
            'file_path': file_path,
            'file_type': file_type,
            'text': None,
            'metadata': None,
            'error': result['error']
        }

    text = result.get('text', '')
    word_count = len(text.split()) if text else 0

    metadata = result.get('metadata', {})
    metadata['word_count'] = word_count
    metadata['char_count'] = len(text)
    metadata['file_size'] = path.stat().st_size

    return {
        'success': True,
        'file_path': file_path,
        'file_type': file_type,
        'text': text,
        'metadata': metadata,
        'error': None
    }

def sanitize_text_for_json(text: str) -> str:
    """Remove or escape control characters that break JSON parsing."""
    if not text:
        return text
    # Remove NULL bytes and other problematic control chars (U+0000-U+001F except \n \r \t)
    # Keep newlines, carriage returns, and tabs as they're valid in JSON strings
    import re
    # Replace control chars except \t \n \r with space
    sanitized = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', text)
    # Normalize multiple spaces to single
    sanitized = re.sub(r' {2,}', ' ', sanitized)
    return sanitized


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(json.dumps({
            'success': False,
            'error': 'Usage: extract-document-text.py <file_path>'
        }))
        sys.exit(1)

    file_path = sys.argv[1]
    result = extract_document(file_path)

    # Sanitize text field to remove control characters
    if result.get('text'):
        result['text'] = sanitize_text_for_json(result['text'])

    print(json.dumps(result, ensure_ascii=False))
